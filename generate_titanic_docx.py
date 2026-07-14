import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell's border using oxml elements
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def apply_apa_table_format(table):
    """
    Apply APA style formatting to word tables:
    - No vertical lines
    - Thin black horizontal lines at the top, bottom of header, and bottom of the table
    - Center alignment
    """
    table.alignment = 1  # Center alignment
    num_rows = len(table.rows)
    
    border_thin = {"sz": 4, "val": "single", "color": "333333", "space": "0"}
    border_none = {"val": "nil"}
    
    for r_idx, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        if r_idx == 0:
            trPr.append(OxmlElement('w:tblHeader'))
            
        for c_idx, cell in enumerate(row.cells):
            borders = {
                "top": border_none,
                "bottom": border_none,
                "left": border_none,
                "right": border_none
            }
            
            if r_idx == 0:
                borders["top"] = border_thin
                borders["bottom"] = border_thin  # Header bottom
            elif r_idx == num_rows - 1:
                borders["bottom"] = border_thin
                
            set_cell_border(cell, **borders)
            
            # Set padding (margins) inside cells
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin in ('top', 'bottom', 'left', 'right'):
                node = OxmlElement(f'w:{margin}')
                node.set(qn('w:w'), '100' if margin in ('top', 'bottom') else '150')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)

def add_page_number(run):
    """
    Insert a dynamic Page Number field into the paragraph run
    """
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def configure_paragraph_style(paragraph, line_spacing=2.0, first_line_indent=0.5, space_before=0, space_after=0):
    """
    Apply standard line spacing, indentations, and margins for APA compliance
    """
    p_format = paragraph.paragraph_format
    p_format.line_spacing = line_spacing
    p_format.first_line_indent = Inches(first_line_indent)
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)

def add_heading(doc, text, level):
    """
    Add headings following APA 7th Edition guidelines:
    - Level 1: Centered, Bold, Title Case
    - Level 2: Left-aligned, Bold, Title Case
    - Level 3: Left-aligned, Bold, Italic, Title Case
    """
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        configure_paragraph_style(p, line_spacing=2.0, first_line_indent=0, space_before=12, space_after=6)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        configure_paragraph_style(p, line_spacing=2.0, first_line_indent=0, space_before=12, space_after=6)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        configure_paragraph_style(p, line_spacing=2.0, first_line_indent=0, space_before=6, space_after=6)
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_body_paragraph(doc, text):
    """
    Add a standard double-spaced body paragraph with 0.5 inches first-line indent
    """
    p = doc.add_paragraph()
    configure_paragraph_style(p, line_spacing=2.0, first_line_indent=0.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_bullet_point(doc, bold_prefix, text):
    """
    Add a bullet point paragraph with standard APA hanging indent style
    """
    p = doc.add_paragraph(style='List Bullet')
    configure_paragraph_style(p, line_spacing=2.0, first_line_indent=0, space_after=4)
    p.paragraph_format.left_indent = Inches(0.5)
    
    run_bold = p.add_run(bold_prefix)
    run_bold.bold = True
    run_bold.font.name = 'Times New Roman'
    run_bold.font.size = Pt(12)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)
    return p

def main():
    doc = Document()
    
    # Configure 1-inch margins on all sides
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Configure Running Header (APA Page Numbering)
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(0)
        hrun = hp.add_run()
        add_page_number(hrun)
        hrun.font.name = 'Times New Roman'
        hrun.font.size = Pt(10)

    # ------------------ TITLE PAGE (APA 7th) ------------------
    for _ in range(3):
        doc.add_paragraph() # Spacing
        
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_paragraph_style(title_p, line_spacing=2.0, first_line_indent=0, space_after=12)
    title_run = title_p.add_run("Modelo Predictivo de Supervivencia en el Siniestro del Titanic mediante Aprendizaje Supervisado de Machine Learning")
    title_run.bold = True
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    
    doc.add_paragraph() # Spacing
    
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_paragraph_style(author_p, line_spacing=2.0, first_line_indent=0)
    
    info_lines = [
        "Dominga Rodríguez",
        "Maestría en Inteligencia Artificial y Aprendizaje Automático",
        "Universidad Francisco Henríquez y Carvajal (UFHEC)",
        "Materia: Aprendizaje Autónomo",
        "Facilitador: [Nombre del Facilitador]",
        "14 de julio de 2026"
    ]
    
    for line in info_lines:
        run = author_p.add_run(line + "\n")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
    doc.add_page_break()

    # ------------------ INTRODUCCIÓN ------------------
    add_heading(doc, "Introducción", level=1)
    add_body_paragraph(doc, 
        "El hundimiento del RMS Titanic en abril de 1912 sigue siendo uno de los desastres marítimos más trágicos e influyentes en la historia moderna. De los 2,224 pasajeros y tripulantes a bordo, más de 1,500 perdieron la vida cuando el barco chocó contra un iceberg en el Atlántico Norte. Esta catástrofe no solo conmocionó al mundo, sino que también impulsó reformas fundamentales en las regulaciones de seguridad marítima internacional, incluyendo la obligatoriedad de botes salvavidas suficientes para todas las personas a bordo."
    )
    add_body_paragraph(doc,
        "Desde una perspectiva analítica y de ciencia de datos, el desastre del Titanic presenta un escenario de estudio único. A pesar de que el naufragio estuvo rodeado de caos y pánico, la supervivencia de los pasajeros no fue completamente aleatoria. Factores estructurales, normas sociales de la época (como el protocolo de \"mujeres y niños primero\"), la ubicación de los camarotes según la clase socioeconómica y la estructura del tamaño de las familias desempeñaron un papel decisivo en la determinación de quién sobrevivió y quién pereció."
    )
    add_body_paragraph(doc,
        "El presente proyecto plantea el diseño y desarrollo de un modelo predictivo basado en Machine Learning Supervisado para estimar la probabilidad de supervivencia de un pasajero en función de sus características socio-demográficas y variables de viaje. Mediante la comparación de múltiples algoritmos supervisados de clasificación, este sistema busca modelar con precisión los patrones de mortalidad del naufragio, sirviendo como una demostración práctica de cómo los métodos predictivos pueden extraer conocimiento valioso de bases de datos históricas estructuradas."
    )

    # ------------------ ANTECEDENTES ------------------
    add_heading(doc, "Antecedentes", level=1)
    add_body_paragraph(doc,
        "El análisis predictivo sobre los datos del Titanic se ha convertido en un punto de referencia (benchmark) fundamental en la literatura de la ciencia de datos y el aprendizaje automático. Diversas investigaciones han explorado cómo variables específicas determinan las probabilidades de supervivencia, validando empíricamente las crónicas históricas del naufragio."
    )
    add_body_paragraph(doc,
        "Estudios clásicos de sociología y economía del comportamiento, como los de Frey et al. (2011), confirman que las normas sociales de caballerosidad (\"mujeres y niños primero\") y los incentivos de supervivencia interactuaron directamente con la clase socioeconómica del pasajero. Sus hallazgos demuestran que los pasajeros de primera clase tuvieron un acceso preferencial a los botes salvavidas, lo que se tradujo en una tasa de supervivencia significativamente superior en comparación con la tercera clase, donde el hacinamiento y la falta de información limitaron la evacuación."
    )
    add_body_paragraph(doc,
        "Por otro lado, la aplicación de algoritmos modernos de clasificación sobre esta base de datos ha evidenciado que los modelos basados en árboles de decisión (como Random Forest y Gradient Boosting) superan a los modelos lineales al capturar interacciones complejas no lineales (por ejemplo, el efecto combinado de ser niño y viajar en tercera clase). Investigadores en computación educativa y analítica predictiva destacan que el dataset del Titanic permite ilustrar de manera transparente el impacto del preprocesamiento de datos (ETL) y el tratamiento de valores faltantes en la precisión final del modelo."
    )

    # ------------------ DEFINICIÓN DEL PROBLEMA ------------------
    add_heading(doc, "Definición del Problema y Pregunta de Investigación", level=1)
    add_heading(doc, "Pregunta de Investigación Orientada a Machine Learning", level=2)
    add_body_paragraph(doc,
        "¿Cómo optimizar la predicción de supervivencia de un pasajero del Titanic mediante un modelo de Machine Learning supervisado de clasificación binaria, utilizando variables demográficas y de viaje como el sexo, la edad, la clase de boleto, la tarifa y el tamaño de la familia?"
    )
    
    add_heading(doc, "Alcance de Predicción del Modelo", level=2)
    add_body_paragraph(doc,
        "El modelo propuesto permitirá:"
    )
    add_bullet_point(doc, "Predecir: ", "Si un pasajero sobrevivió (1) o no (0) al naufragio basándose en su perfil de datos.")
    add_bullet_point(doc, "Clasificar: ", "A los pasajeros en grupos de alta o baja probabilidad de supervivencia para auditar las disparidades del evento.")
    add_bullet_point(doc, "Identificar: ", "Qué variables (como la clase o el sexo) tuvieron el mayor peso o importancia predictiva en el desenlace del siniestro.")
    
    add_heading(doc, "¿A quién impacta?", level=2)
    add_body_paragraph(doc,
        "Este proyecto tiene un fuerte carácter didáctico y académico, impactando directamente a estudiantes y docentes de la Maestría en Inteligencia Artificial y Aprendizaje Automático en la Universidad Francisco Henríquez y Carvajal (UFHEC), al proporcionar un pipeline transparente y reproducible bajo la metodología de desarrollo de ciclo de vida CRISP-ML(Q)."
    )

    # ------------------ ARQUITECTURA DEL MODELO ------------------
    add_heading(doc, "Arquitectura del Modelo de Machine Learning y Decisiones", level=1)
    add_heading(doc, "Variable Objetivo (Target)", level=2)
    add_body_paragraph(doc,
        "Survived (Clasificación Binaria): 1 = Sobrevivió, 0 = No sobrevivió (Falleció)."
    )
    
    add_heading(doc, "Variables Independientes (Features)", level=2)
    add_body_paragraph(doc,
        "Las variables independientes corresponden a las características del pasajero y su boleto:"
    )
    add_bullet_point(doc, "Pclass: ", "Clase de Pasajero (1.ª, 2.ª o 3.ª clase). Proxy socioeconómico.")
    add_bullet_point(doc, "Sex: ", "Género del pasajero (male/female).")
    add_bullet_point(doc, "Age: ", "Edad del pasajero en años.")
    add_bullet_point(doc, "Fare: ", "Tarifa pagada por el boleto.")
    add_bullet_point(doc, "FamilySize: ", "Tamaño de la familia a bordo (SibSp + Parch + 1).")
    add_bullet_point(doc, "IsAlone: ", "Indicador de si viaja solo (1) o acompañado (0).")
    add_bullet_point(doc, "Embarked: ", "Puerto de Embarque (Cherbourg, Queenstown, Southampton).")

    # --- TABLA 1: ARQUITECTURA ---
    doc.add_paragraph()
    table_title_1 = doc.add_paragraph()
    configure_paragraph_style(table_title_1, line_spacing=1.15, first_line_indent=0, space_after=2)
    run_t1_1 = table_title_1.add_run("Tabla 1\n")
    run_t1_1.bold = True
    run_t1_1.font.name = 'Times New Roman'
    run_t1_2 = table_title_1.add_run("Arquitectura de Datos del Sistema y Decisiones Tecnológicas")
    run_t1_2.italic = True
    run_t1_2.font.name = 'Times New Roman'
    
    table1_data = [
        ("Componente", "Descripción"),
        ("Tipo de IA", "Machine Learning Supervisado. El modelo aprende a partir de etiquetas reales de supervivencia."),
        ("Problema", "Clasificación binaria (Clases: 1 / 0)."),
        ("Algoritmo Principal", "Random Forest Classifier (para interacciones complejas no lineales)."),
        ("Algoritmo Base", "Regresión Logística (para coeficientes matemáticos e inferencia en frontend JS)."),
        ("Variables", "Pclass, Sex, Age, SibSp, Parch, Fare, Embarked, FamilySize, IsAlone."),
        ("Métricas de Evaluación", "Accuracy, Precision, Recall, F1-Score y AUC-ROC."),
        ("Herramientas", "Python, Scikit-learn, Pandas, NumPy, Streamlit, HTML5/CSS3/JavaScript.")
    ]
    
    t1 = doc.add_table(rows=len(table1_data), cols=2)
    for r_idx, row in enumerate(table1_data):
        for c_idx, val in enumerate(row):
            cell = t1.cell(r_idx, c_idx)
            cell.text = val
            p = cell.paragraphs[0]
            configure_paragraph_style(p, line_spacing=1.15, first_line_indent=0)
            if r_idx == 0:
                p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(10)
            
    apply_apa_table_format(t1)

    add_heading(doc, "Justificación del Modelo Seleccionado", level=2)
    add_body_paragraph(doc,
        "Se seleccionan Random Forest y Regresión Logística como los pilares del proyecto. El Random Forest es ideal para capturar las complejas interacciones no lineales de las variables demográficas en el barco (por ejemplo, el efecto del sexo combinado con la clase socioeconómica). La Regresión Logística, al proporcionar coeficientes numéricos explícitos y un cálculo de probabilidad directo mediante una sigmoide lineal, resulta óptima para el despliegue ligero en JavaScript cliente."
    )

    # ------------------ DATOS NECESARIOS ------------------
    add_heading(doc, "Datos Necesarios para el Desarrollo del Modelo", level=1)
    
    # --- TABLA 2: DATOS REQUERIDOS ---
    doc.add_paragraph()
    table_title_2 = doc.add_paragraph()
    configure_paragraph_style(table_title_2, line_spacing=1.15, first_line_indent=0, space_after=2)
    run_t2_1 = table_title_2.add_run("Tabla 2\n")
    run_t2_1.bold = True
    run_t2_1.font.name = 'Times New Roman'
    run_t2_2 = table_title_2.add_run("Tipos de Datos Requeridos y Fuentes de Obtención")
    run_t2_2.italic = True
    run_t2_2.font.name = 'Times New Roman'
    
    table2_data = [
        ("Tipo de dato", "Descripción", "Fuente de obtención", "Tecnología / Formato"),
        ("Datos Demográficos", "Género, edad y estatus familiar del pasajero.", "Dataset histórico del Titanic.", "CSV (titanic.csv)"),
        ("Datos de Viaje", "Clase de cabina, puerto de embarque, tarifa y código de boleto.", "Registros de la lista de pasajeros.", "CSV / Pandas DataFrame")
    ]
    
    t2 = doc.add_table(rows=len(table2_data), cols=4)
    for r_idx, row in enumerate(table2_data):
        for c_idx, val in enumerate(row):
            cell = t2.cell(r_idx, c_idx)
            cell.text = val
            p = cell.paragraphs[0]
            configure_paragraph_style(p, line_spacing=1.15, first_line_indent=0)
            if r_idx == 0:
                p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(10)
            
    apply_apa_table_format(t2)

    # ------------------ PROCESO ETL Y EDA ------------------
    add_heading(doc, "Datos y Analítica para la Toma de Decisiones", level=1)
    add_heading(doc, "Proceso ETL (Extract, Transform, Load)", level=2)
    add_body_paragraph(doc,
        "Extract: Extracción del manifiesto estructurado de pasajeros desde titanic.csv.\n"
        "Transform: Imputación de edad mediante la mediana de grupos específicos (Pclass y Sex); imputación de embarque mediante moda. Creación de variables de ingeniería: FamilySize e IsAlone. Codificación binaria para Sex (0/1) y estandarización para variables continuas (Age, Fare).\n"
        "Load: División estructurada del dataset en subconjuntos de entrenamiento (70%) y prueba (30%)."
    )
    
    add_heading(doc, "Análisis Exploratorio de Datos (EDA)", level=2)
    add_body_paragraph(doc,
        "El EDA visual y estadístico confirmó la alta tasa de supervivencia en mujeres (74%) contra hombres (19%), así como el gradiente socioeconómico en la supervivencia (1.ª clase: 63%, 2.ª clase: 47%, 3.ª clase: 24%)."
    )

    # ------------------ DESAFÍOS DE GESTIÓN DE DATOS ------------------
    add_heading(doc, "Planificación y Desafíos de la Gestión de Datos", level=1)
    
    # --- TABLA 3: DESAFÍOS ---
    doc.add_paragraph()
    table_title_3 = doc.add_paragraph()
    configure_paragraph_style(table_title_3, line_spacing=1.15, first_line_indent=0, space_after=2)
    run_t3_1 = table_title_3.add_run("Tabla 3\n")
    run_t3_1.bold = True
    run_t3_1.font.name = 'Times New Roman'
    run_t3_2 = table_title_3.add_run("Desafíos Estratégicos de los Datos")
    run_t3_2.italic = True
    run_t3_2.font.name = 'Times New Roman'
    
    table3_data = [
        ("Factor", "Fortalezas", "Oportunidades", "Debilidades", "Amenazas"),
        ("Calidad de datos", "Formato tabular limpio.", "Aplicar técnicas de imputación avanzadas.", "Faltantes sustanciales en Cabin.", "Sesgo por imputación incorrecta de Age."),
        ("Sesgo de datos", "Datos reales del hecho.", "Analizar disparidades socio-demográficas históricas.", "Desbalance moderado (61% fallecidos).", "El modelo lineal puede ignorar factores no dominantes.")
    ]
    
    t3 = doc.add_table(rows=len(table3_data), cols=5)
    for r_idx, row in enumerate(table3_data):
        for c_idx, val in enumerate(row):
            cell = t3.cell(r_idx, c_idx)
            cell.text = val
            p = cell.paragraphs[0]
            configure_paragraph_style(p, line_spacing=1.15, first_line_indent=0)
            if r_idx == 0:
                p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(9.5)
            
    apply_apa_table_format(t3)

    # ------------------ MATRIZ DOFA ------------------
    doc.add_paragraph()
    table_title_4 = doc.add_paragraph()
    configure_paragraph_style(table_title_4, line_spacing=1.15, first_line_indent=0, space_after=2)
    run_t4_1 = table_title_4.add_run("Tabla 4\n")
    run_t4_1.bold = True
    run_t4_1.font.name = 'Times New Roman'
    run_t4_2 = table_title_4.add_run("Matriz DOFA – Gestión de Datos en el Proyecto Titanic")
    run_t4_2.italic = True
    run_t4_2.font.name = 'Times New Roman'
    
    table4_data = [
        ("Fortalezas (F)", "Debilidades (D)", "Oportunidades (O)", "Amenazas (A)"),
        ("Registros tabulares bien estructurados y oficiales.", "Datos históricos limitados que impiden recolectar nuevas muestras.", "Uso académico para ilustrar sesgos sociales en algoritmos.", "El sesgo de género inherente puede nublar la influencia de otras variables.")
    ]
    
    t4 = doc.add_table(rows=len(table4_data), cols=4)
    for r_idx, row in enumerate(table4_data):
        for c_idx, val in enumerate(row):
            cell = t4.cell(r_idx, c_idx)
            cell.text = val
            p = cell.paragraphs[0]
            configure_paragraph_style(p, line_spacing=1.15, first_line_indent=0)
            if r_idx == 0:
                p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(10)
            
    apply_apa_table_format(t4)

    # ------------------ ETICA Y SMART ------------------
    add_heading(doc, "Objetivos y Dimensión Ética", level=1)
    add_heading(doc, "Objetivos SMART", level=2)
    add_bullet_point(doc, "Específico: ", "Modelar la supervivencia binaria en el Titanic.")
    add_bullet_point(doc, "Medible: ", "Accuracy superior al 80% en validación de test.")
    add_bullet_point(doc, "Alcanzable: ", "Mediante el uso de Scikit-learn sobre el dataset original de 891 registros.")
    add_bullet_point(doc, "Relevante: ", "Ilustra el proceso completo de CRISP-ML(Q).")
    add_bullet_point(doc, "Temporal: ", "Desarrollado y finalizado dentro del cronograma lectivo.")

    # --- TABLA 5: CONSIDERACIONES ETICAS ---
    doc.add_paragraph()
    table_title_5 = doc.add_paragraph()
    configure_paragraph_style(table_title_5, line_spacing=1.15, first_line_indent=0, space_after=2)
    run_t5_1 = table_title_5.add_run("Tabla 5\n")
    run_t5_1.bold = True
    run_t5_1.font.name = 'Times New Roman'
    run_t5_2 = table_title_5.add_run("Consideraciones Éticas de la Plataforma")
    run_t5_2.italic = True
    run_t5_2.font.name = 'Times New Roman'
    
    table5_data = [
        ("Consideración ética", "Enfoque de Mitigación", "¿Aplica?"),
        ("Dignidad e Identidad", "Los nombres de los pasajeros se omiten del entrenamiento para evitar estigmatización y centrar el análisis en perfiles socio-demográficos.", "Sí"),
        ("Transparencia Algorítmica", "Las reglas complejas del Random Forest se explican mediante la interpretabilidad de un Árbol de Decisión complementario.", "Sí"),
        ("Sesgo Histórico", "Se expone explícitamente el sesgo de clase y género del naufragio para evitar replicarlo en la interpretación moderna.", "Sí")
    ]
    
    t5 = doc.add_table(rows=len(table5_data), cols=3)
    for r_idx, row in enumerate(table5_data):
        for c_idx, val in enumerate(row):
            cell = t5.cell(r_idx, c_idx)
            cell.text = val
            p = cell.paragraphs[0]
            configure_paragraph_style(p, line_spacing=1.15, first_line_indent=0)
            if r_idx == 0:
                p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(10)
            
    apply_apa_table_format(t5)

    # ------------------ ELEVATOR PITCH ------------------
    add_heading(doc, "Guía para el Elevator Pitch (5 minutos)", level=1)
    
    # --- TABLA 6: ELEVATOR PITCH ---
    doc.add_paragraph()
    table_title_6 = doc.add_paragraph()
    configure_paragraph_style(table_title_6, line_spacing=1.15, first_line_indent=0, space_after=2)
    run_t6_1 = table_title_6.add_run("Tabla 6\n")
    run_t6_1.bold = True
    run_t6_1.font.name = 'Times New Roman'
    run_t6_2 = table_title_6.add_run("Estructura de la Presentación de 5 Minutos")
    run_t6_2.italic = True
    run_t6_2.font.name = 'Times New Roman'
    
    table6_data = [
        ("Tiempo", "Contenido de la presentación"),
        ("Minuto 0–1: El problema", "Presentación del naufragio y justificación de que la supervivencia no fue al azar."),
        ("Minuto 1–2: La solución", "Uso de clasificación binaria (Random Forest y Regresión Logística) para modelar la supervivencia."),
        ("Minuto 2–3: Los datos", "Tratamiento de nulos de edad, agregación del tamaño familiar e importancia de la tarifa."),
        ("Minuto 3–4: La ética", "Transparencia del algoritmo y análisis del sesgo de clase y género histórico."),
        ("Minuto 4–5: El impacto", "Visualización de las aplicaciones web e impacto pedagógico en UFHEC.")
    ]
    
    t6 = doc.add_table(rows=len(table6_data), cols=2)
    for r_idx, row in enumerate(table6_data):
        for c_idx, val in enumerate(row):
            cell = t6.cell(r_idx, c_idx)
            cell.text = val
            p = cell.paragraphs[0]
            configure_paragraph_style(p, line_spacing=1.15, first_line_indent=0)
            if r_idx == 0:
                p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(10)
            
    apply_apa_table_format(t6)

    # ------------------ ANALISIS CRITICO ------------------
    add_heading(doc, "Análisis Crítico y Estratégico del Equipo", level=1)
    add_heading(doc, "¿Cuál es el mayor desafío en su proyecto?", level=2)
    add_body_paragraph(doc,
        "El mayor desafío es el sesgo histórico y el desbalance moderado. Dado que el protocolo \"mujeres y niños primero\" se aplicó con rigidez en las cubiertas superiores, la variable Sex tiene una importancia predictiva tan alta que los clasificadores simples tienden a ignorar las demás variables. Un hombre joven de 1.ª clase que pagó una tarifa muy alta y tenía acceso preferencial a los botes podría ser erróneamente clasificado como \"No sobreviviente\" debido al peso abrumador del factor de género en el modelo lineal."
    )
    add_heading(doc, "¿Cómo lo mitigarían estratégicamente?", level=2)
    add_body_paragraph(doc,
        "Para mitigar esto, implementamos modelos de ensamble no lineales como Random Forest y ajustamos los pesos de clase (class_weight='balanced'). Esto obliga al optimizador a prestar atención a las combinaciones complejas de variables (por ejemplo, pasajeros masculinos de primera clase que viajaban acompañados) y reduce la tasa de falsos negativos en hombres con condiciones socioeconómicas favorables para la supervivencia."
    )

    # ------------------ REFERENCIAS ------------------
    doc.add_page_break()
    add_heading(doc, "Referencias", level=1)
    
    refs = [
        "Frey, B. S., Savage, D. A., & Torgler, B. (2011). Behavior under extreme conditions: The Titanic disaster. Journal of Economic Perspectives, 25(1), 209–222. https://doi.org/10.1257/jep.25.1.209",
        "Geron, A. (2019). Hands-on Machine Learning with Scikit-Learn, Keras, and TensorFlow (2nd ed.). O'Reilly Media.",
        "Hall, P., & Gill, N. (2019). An introduction to machine learning interpretability (2nd ed.). O'Reilly Media."
    ]
    
    for ref in refs:
        p = doc.add_paragraph()
        configure_paragraph_style(p, line_spacing=2.0, first_line_indent=0)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    doc.add_page_break()

    # ------------------ ANEXOS ------------------
    add_heading(doc, "Anexos", level=1)
    add_heading(doc, "Anexo 1: Datasets y Enlaces de Interés", level=2)
    
    # --- TABLA 7: DATASETS ---
    doc.add_paragraph()
    table_title_7 = doc.add_paragraph()
    configure_paragraph_style(table_title_7, line_spacing=1.15, first_line_indent=0, space_after=2)
    run_t7_1 = table_title_7.add_run("Tabla 7\n")
    run_t7_1.bold = True
    run_t7_1.font.name = 'Times New Roman'
    run_t7_2 = table_title_7.add_run("Datasets Recomendados y Enlaces de Kaggle")
    run_t7_2.italic = True
    run_t7_2.font.name = 'Times New Roman'
    
    table7_data = [
        ("Nombre del dataset", "Descripción", "Variables útiles", "Enlace"),
        ("Titanic: Machine Learning from Disaster", "Dataset oficial de Kaggle para modelar la supervivencia en el naufragio.", "Survived, Pclass, Sex, Age, SibSp, Parch, Fare, Cabin, Embarked.", "Kaggle"),
        ("Titanic passenger list", "Lista completa de pasajeros históricos del Titanic con detalles adicionales.", "Passenger records, details, class.", "Kaggle")
    ]
    
    t7 = doc.add_table(rows=len(table7_data), cols=4)
    for r_idx, row in enumerate(table7_data):
        for c_idx, val in enumerate(row):
            cell = t7.cell(r_idx, c_idx)
            cell.text = val
            p = cell.paragraphs[0]
            configure_paragraph_style(p, line_spacing=1.15, first_line_indent=0)
            if r_idx == 0:
                p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(9.5)
            
    apply_apa_table_format(t7)

    output_paths = ["Documento_Titanic.docx"]
    for path in output_paths:
        doc.save(path)
        print(f"Documento guardado exitosamente en: {os.path.abspath(path)}")

if __name__ == "__main__":
    main()
